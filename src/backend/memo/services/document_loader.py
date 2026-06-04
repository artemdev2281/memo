import hashlib
import os
from dataclasses import dataclass

SUPPORTED: frozenset[str] = frozenset({".pdf", ".docx", ".md", ".txt"})


@dataclass
class Document:
    path: str
    text: str
    file_hash: str


def compute_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def load(path: str) -> Document:
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED:
        raise ValueError(f"Unsupported format: {ext!r}")

    h = compute_hash(path)

    if ext == ".pdf":
        import fitz  # PyMuPDF

        doc = fitz.open(path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
    elif ext == ".docx":
        from docx import Document as DocxDocument

        d = DocxDocument(path)
        text = "\n".join(p.text for p in d.paragraphs if p.text)
    else:
        with open(path, encoding="utf-8", errors="replace") as f:
            text = f.read()

    return Document(path=path, text=text, file_hash=h)
